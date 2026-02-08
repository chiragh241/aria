"""RAG (Retrieval Augmented Generation) pipeline."""

from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import TYPE_CHECKING, Any

from ..utils.config import get_settings
from ..utils.logging import get_logger
from .long_term import LongTermMemory, SearchResult
from .cognee_graph import CogneeGraphMemory

logger = get_logger(__name__)

if TYPE_CHECKING:
    from .user_profile import UserProfile


@dataclass
class RAGResult:
    """Result from a RAG query."""

    query: str
    context: str
    sources: list[dict[str, Any]]
    relevance_score: float


@dataclass
class Document:
    """A document for ingestion."""

    content: str
    source: str
    metadata: dict[str, Any] = field(default_factory=dict)
    chunk_size: int = 500
    chunk_overlap: int = 50


class RAGPipeline:
    """
    RAG pipeline for context-aware responses.

    Features:
    - Document ingestion with chunking
    - Semantic search
    - Context building for LLM queries
    - Source attribution
    """

    def __init__(
        self,
        memory: LongTermMemory | None = None,
        cognee_memory: CogneeGraphMemory | None = None,
    ) -> None:
        self.settings = get_settings()
        self._memory = memory or LongTermMemory()
        self._cognee = cognee_memory or CogneeGraphMemory()
        self._initialized = False

    async def initialize(self) -> None:
        """Initialize the RAG pipeline."""
        if self._initialized:
            return

        await self._memory.initialize()
        await self._cognee.initialize()
        self._initialized = True
        logger.info(
            "RAG pipeline initialized",
            knowledge_graph=self._cognee.is_available(),
        )

    def _chunk_text(
        self,
        text: str,
        chunk_size: int = 500,
        chunk_overlap: int = 50,
    ) -> list[str]:
        """
        Split text into overlapping chunks.

        Args:
            text: Text to chunk
            chunk_size: Target chunk size in characters
            chunk_overlap: Overlap between chunks

        Returns:
            List of text chunks
        """
        if len(text) <= chunk_size:
            return [text]

        chunks = []
        start = 0

        while start < len(text):
            end = start + chunk_size

            # Try to break at sentence boundary
            if end < len(text):
                # Look for sentence end
                for sep in [". ", ".\n", "! ", "? ", "\n\n"]:
                    last_sep = text[start:end].rfind(sep)
                    if last_sep > chunk_size // 2:
                        end = start + last_sep + len(sep)
                        break

            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)

            start = end - chunk_overlap

        return chunks

    async def ingest_document(
        self,
        document: Document,
    ) -> int:
        """
        Ingest a document into the RAG system.

        Args:
            document: Document to ingest

        Returns:
            Number of chunks created
        """
        if not self._initialized:
            await self.initialize()

        chunks = self._chunk_text(
            document.content,
            chunk_size=document.chunk_size,
            chunk_overlap=document.chunk_overlap,
        )

        for i, chunk in enumerate(chunks):
            metadata = {
                **document.metadata,
                "source": document.source,
                "chunk_index": i,
                "total_chunks": len(chunks),
            }

            await self._memory.add(
                content=chunk,
                metadata=metadata,
                source=document.source,
            )

        # Also feed into knowledge graph for entity/relationship extraction
        if self._cognee.is_available():
            await self._cognee.add_knowledge(document.content)
            # Auto-process if enabled (runs in background to avoid blocking)
            if self.settings.memory.knowledge_graph.auto_process_after_ingest:
                import asyncio
                asyncio.create_task(self._auto_process_knowledge())

        logger.debug(
            "Ingested document",
            source=document.source,
            chunks=len(chunks),
        )

        return len(chunks)

    async def _auto_process_knowledge(self) -> None:
        """Background task to process knowledge graph after ingestion."""
        try:
            await self.process_knowledge_graph()
            logger.info("Knowledge graph auto-processed after ingestion")
        except Exception as e:
            logger.warning("Knowledge graph auto-process failed", error=str(e))

    async def ingest_file(
        self,
        file_path: str | Path,
        metadata: dict[str, Any] | None = None,
    ) -> int:
        """
        Ingest a file into the RAG system.

        Args:
            file_path: Path to the file
            metadata: Optional metadata

        Returns:
            Number of chunks created
        """
        path = Path(file_path).expanduser()

        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")

        # Read file content based on type
        suffix = path.suffix.lower()

        if suffix in [".txt", ".md", ".py", ".js", ".json", ".yaml", ".yml"]:
            content = path.read_text(encoding="utf-8", errors="replace")
        elif suffix == ".pdf":
            content = await self._extract_pdf(path)
        elif suffix in [".docx"]:
            content = await self._extract_docx(path)
        else:
            # Try to read as text
            content = path.read_text(encoding="utf-8", errors="replace")

        document = Document(
            content=content,
            source=str(path),
            metadata={
                "filename": path.name,
                "file_type": suffix,
                **(metadata or {}),
            },
        )

        return await self.ingest_document(document)

    async def _extract_pdf(self, path: Path) -> str:
        """Extract text from a PDF file."""
        try:
            from pypdf import PdfReader

            reader = PdfReader(path)
            text_parts = []
            for page in reader.pages:
                text_parts.append(page.extract_text())
            return "\n\n".join(text_parts)
        except ImportError:
            logger.warning("pypdf not installed, cannot extract PDF")
            return ""

    async def _extract_docx(self, path: Path) -> str:
        """Extract text from a DOCX file."""
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(path)
            text_parts = []
            for para in doc.paragraphs:
                text_parts.append(para.text)
            return "\n\n".join(text_parts)
        except ImportError:
            logger.warning("python-docx not installed, cannot extract DOCX")
            return ""

    async def ingest_conversation(
        self,
        messages: list[dict[str, str]],
        conversation_id: str,
    ) -> int:
        """
        Ingest a conversation into RAG for future reference.

        Args:
            messages: List of messages with 'role' and 'content'
            conversation_id: Unique conversation identifier

        Returns:
            Number of chunks created
        """
        # Format conversation as text
        parts = []
        for msg in messages:
            role = msg.get("role", "unknown").upper()
            content = msg.get("content", "")
            parts.append(f"{role}: {content}")

        content = "\n\n".join(parts)

        document = Document(
            content=content,
            source=f"conversation:{conversation_id}",
            metadata={
                "type": "conversation",
                "conversation_id": conversation_id,
                "message_count": len(messages),
            },
        )

        return await self.ingest_document(document)

    async def query(
        self,
        query: str,
        top_k: int = 5,
        min_relevance: float = 0.3,
        source_filter: str | None = None,
    ) -> list[dict[str, Any]]:
        """
        Query the RAG system.

        Args:
            query: Search query
            top_k: Maximum results
            min_relevance: Minimum relevance score (0-1)
            source_filter: Optional source prefix filter

        Returns:
            List of relevant context items
        """
        if not self._initialized:
            await self.initialize()

        where = None
        if source_filter:
            where = {"source": {"$contains": source_filter}}

        results = await self._memory.search(
            query=query,
            top_k=top_k,
            where=where,
        )

        # Filter by relevance
        relevant = [
            {
                "content": r.document.content,
                "source": r.document.metadata.get("source", "unknown"),
                "score": r.score,
                "metadata": r.document.metadata,
            }
            for r in results
            if r.score >= min_relevance
        ]

        # Merge with knowledge graph results for richer context
        if self._cognee.is_available():
            graph_results = await self._cognee.search(query, top_k=max(2, top_k // 2))
            for gr in graph_results:
                if gr.get("score", 0) >= min_relevance:
                    relevant.append(gr)

        # Sort by score descending and deduplicate
        relevant.sort(key=lambda x: x.get("score", 0), reverse=True)
        return relevant[:top_k]

    async def get_context(
        self,
        query: str,
        max_tokens: int = 2000,
        top_k: int = 5,
    ) -> RAGResult:
        """
        Get context for a query to augment LLM responses.

        Args:
            query: The user's query
            max_tokens: Maximum tokens for context
            top_k: Maximum sources to consider

        Returns:
            RAGResult with context and sources
        """
        results = await self.query(query, top_k=top_k)

        if not results:
            return RAGResult(
                query=query,
                context="",
                sources=[],
                relevance_score=0,
            )

        # Build context, respecting token limit
        context_parts = []
        sources = []
        total_chars = 0
        max_chars = max_tokens * 4  # Rough estimate

        for item in results:
            content = item["content"]
            if total_chars + len(content) > max_chars:
                # Truncate if needed
                remaining = max_chars - total_chars
                if remaining > 100:
                    content = content[:remaining] + "..."
                else:
                    break

            context_parts.append(content)
            sources.append({
                "source": item["source"],
                "score": item["score"],
            })
            total_chars += len(content)

        context = "\n\n---\n\n".join(context_parts)
        avg_score = sum(s["score"] for s in sources) / len(sources) if sources else 0

        return RAGResult(
            query=query,
            context=context,
            sources=sources,
            relevance_score=avg_score,
        )

    async def delete_source(self, source: str) -> int:
        """
        Delete all documents from a source.

        Args:
            source: Source identifier

        Returns:
            Number of documents deleted
        """
        return await self._memory.delete_many(
            where={"source": source}
        )

    async def add_user_profile_to_knowledge(
        self,
        user_id: str,
        profile: "UserProfile",
    ) -> bool:
        """
        Sync user profile into the knowledge graph.
        Keeps Cognee in sync with user profiles so graph search can find
        preferences, relationships, and facts alongside ingested documents.

        Args:
            user_id: User identifier
            profile: User profile object

        Returns:
            True if content was added
        """
        if not self._cognee.is_available():
            return False
        return await self._cognee.add_user_profile_knowledge(user_id, profile)

    async def process_knowledge_graph(self) -> bool:
        """Trigger cognee to process accumulated knowledge into the graph."""
        if not self._cognee.is_available():
            return False
        return await self._cognee.process_knowledge()

    def get_stats(self) -> dict[str, Any]:
        """Get RAG pipeline statistics."""
        return {
            "initialized": self._initialized,
            "memory_stats": self._memory.get_stats(),
            "knowledge_graph": self._cognee.get_stats(),
        }
