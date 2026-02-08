"""Document processing skill for PDF and DOCX."""

from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from ..base import BaseSkill, SkillResult


class DocumentsSkill(BaseSkill):
    """
    Skill for document creation and processing.

    Supports PDF and DOCX creation, reading, and manipulation.
    """

    name = "documents"
    description = "Document creation and processing"
    version = "1.0.0"

    def _register_capabilities(self) -> None:
        """Register document capabilities."""
        self.register_capability(
            name="create_pdf",
            description="Create a PDF document",
            parameters={
                "type": "object",
                "properties": {
                    "output_path": {"type": "string", "description": "Output PDF path"},
                    "content": {"type": "string", "description": "Text content"},
                    "title": {"type": "string", "description": "Document title"},
                    "author": {"type": "string", "description": "Document author"},
                },
                "required": ["output_path", "content"],
            },
            security_action="write_files",
        )

        self.register_capability(
            name="create_docx",
            description="Create a Word document",
            parameters={
                "type": "object",
                "properties": {
                    "output_path": {"type": "string", "description": "Output DOCX path"},
                    "content": {"type": "string", "description": "Text content (supports markdown)"},
                    "title": {"type": "string", "description": "Document title"},
                },
                "required": ["output_path", "content"],
            },
            security_action="write_files",
        )

        self.register_capability(
            name="read_pdf",
            description="Extract text from PDF",
            parameters={
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "PDF file path"},
                    "pages": {"type": "array", "items": {"type": "integer"}, "description": "Specific pages to read"},
                },
                "required": ["path"],
            },
            security_action="read_files",
        )

        self.register_capability(
            name="read_docx",
            description="Extract text from Word document",
            parameters={
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "DOCX file path"},
                },
                "required": ["path"],
            },
            security_action="read_files",
        )

        self.register_capability(
            name="merge_pdfs",
            description="Merge multiple PDFs into one",
            parameters={
                "type": "object",
                "properties": {
                    "input_paths": {"type": "array", "items": {"type": "string"}, "description": "Input PDF paths"},
                    "output_path": {"type": "string", "description": "Output PDF path"},
                },
                "required": ["input_paths", "output_path"],
            },
            security_action="write_files",
        )

        self.register_capability(
            name="pdf_info",
            description="Get PDF metadata and info",
            parameters={
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "PDF file path"},
                },
                "required": ["path"],
            },
            security_action="read_files",
        )

    async def execute(self, capability: str, **kwargs: Any) -> SkillResult:
        """Execute a document capability."""
        start_time = datetime.now(timezone.utc)

        handlers = {
            "create_pdf": self._create_pdf,
            "create_docx": self._create_docx,
            "read_pdf": self._read_pdf,
            "read_docx": self._read_docx,
            "merge_pdfs": self._merge_pdfs,
            "pdf_info": self._pdf_info,
        }

        handler = handlers.get(capability)
        if not handler:
            return self._error_result(f"Unknown capability: {capability}", start_time)

        try:
            result = await handler(**kwargs)
            return self._success_result(result, start_time)
        except ImportError as e:
            return self._error_result(
                f"Required library not installed: {str(e)}. "
                "Run: pip install reportlab python-docx pypdf",
                start_time,
            )
        except Exception as e:
            return self._error_result(str(e), start_time)

    async def _create_pdf(
        self,
        output_path: str,
        content: str,
        title: str | None = None,
        author: str | None = None,
    ) -> dict[str, Any]:
        """Create a PDF document."""
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

        output_file = Path(output_path).expanduser()
        output_file.parent.mkdir(parents=True, exist_ok=True)

        doc = SimpleDocTemplate(
            str(output_file),
            pagesize=letter,
            title=title or "Document",
            author=author or "Aria",
        )

        styles = getSampleStyleSheet()
        story = []

        # Add title if provided
        if title:
            title_style = ParagraphStyle(
                "Title",
                parent=styles["Heading1"],
                fontSize=24,
                spaceAfter=30,
            )
            story.append(Paragraph(title, title_style))

        # Add content paragraphs
        for paragraph in content.split("\n\n"):
            if paragraph.strip():
                # Handle basic markdown
                if paragraph.startswith("# "):
                    story.append(Paragraph(paragraph[2:], styles["Heading1"]))
                elif paragraph.startswith("## "):
                    story.append(Paragraph(paragraph[3:], styles["Heading2"]))
                elif paragraph.startswith("### "):
                    story.append(Paragraph(paragraph[4:], styles["Heading3"]))
                else:
                    story.append(Paragraph(paragraph, styles["Normal"]))
                story.append(Spacer(1, 12))

        doc.build(story)

        return {
            "output_path": str(output_file),
            "file_size": output_file.stat().st_size,
            "title": title,
        }

    async def _create_docx(
        self,
        output_path: str,
        content: str,
        title: str | None = None,
    ) -> dict[str, Any]:
        """Create a Word document."""
        from docx import Document
        from docx.shared import Pt

        output_file = Path(output_path).expanduser()
        output_file.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()

        # Add title if provided
        if title:
            heading = doc.add_heading(title, level=0)

        # Add content paragraphs
        for paragraph in content.split("\n\n"):
            if paragraph.strip():
                # Handle basic markdown
                if paragraph.startswith("# "):
                    doc.add_heading(paragraph[2:], level=1)
                elif paragraph.startswith("## "):
                    doc.add_heading(paragraph[3:], level=2)
                elif paragraph.startswith("### "):
                    doc.add_heading(paragraph[4:], level=3)
                elif paragraph.startswith("- "):
                    # Bullet list
                    for line in paragraph.split("\n"):
                        if line.startswith("- "):
                            doc.add_paragraph(line[2:], style="List Bullet")
                else:
                    doc.add_paragraph(paragraph)

        doc.save(str(output_file))

        return {
            "output_path": str(output_file),
            "file_size": output_file.stat().st_size,
            "title": title,
        }

    async def _read_pdf(
        self,
        path: str,
        pages: list[int] | None = None,
    ) -> dict[str, Any]:
        """Read text from PDF."""
        from pypdf import PdfReader

        pdf_path = Path(path).expanduser()
        reader = PdfReader(pdf_path)

        text_parts = []
        page_indices = pages if pages else range(len(reader.pages))

        for i in page_indices:
            if 0 <= i < len(reader.pages):
                page = reader.pages[i]
                text = page.extract_text()
                text_parts.append({
                    "page": i + 1,
                    "text": text,
                })

        return {
            "path": str(pdf_path),
            "total_pages": len(reader.pages),
            "pages_read": len(text_parts),
            "content": text_parts,
            "full_text": "\n\n".join(p["text"] for p in text_parts),
        }

    async def _read_docx(self, path: str) -> dict[str, Any]:
        """Read text from Word document."""
        from docx import Document

        docx_path = Path(path).expanduser()
        doc = Document(docx_path)

        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append({
                    "text": para.text,
                    "style": para.style.name if para.style else None,
                })

        # Extract tables
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)

        full_text = "\n\n".join(p["text"] for p in paragraphs)

        return {
            "path": str(docx_path),
            "paragraphs": len(paragraphs),
            "tables": len(tables),
            "content": paragraphs,
            "table_data": tables,
            "full_text": full_text,
        }

    async def _merge_pdfs(
        self,
        input_paths: list[str],
        output_path: str,
    ) -> dict[str, Any]:
        """Merge multiple PDFs."""
        from pypdf import PdfMerger

        output_file = Path(output_path).expanduser()
        output_file.parent.mkdir(parents=True, exist_ok=True)

        merger = PdfMerger()

        for path in input_paths:
            pdf_path = Path(path).expanduser()
            if pdf_path.exists():
                merger.append(str(pdf_path))

        merger.write(str(output_file))
        merger.close()

        return {
            "output_path": str(output_file),
            "input_count": len(input_paths),
            "file_size": output_file.stat().st_size,
        }

    async def _pdf_info(self, path: str) -> dict[str, Any]:
        """Get PDF metadata."""
        from pypdf import PdfReader

        pdf_path = Path(path).expanduser()
        reader = PdfReader(pdf_path)

        metadata = reader.metadata or {}

        return {
            "path": str(pdf_path),
            "pages": len(reader.pages),
            "file_size": pdf_path.stat().st_size,
            "encrypted": reader.is_encrypted,
            "metadata": {
                "title": metadata.get("/Title"),
                "author": metadata.get("/Author"),
                "subject": metadata.get("/Subject"),
                "creator": metadata.get("/Creator"),
                "producer": metadata.get("/Producer"),
                "creation_date": str(metadata.get("/CreationDate")),
                "modification_date": str(metadata.get("/ModDate")),
            },
        }
